from openai import OpenAI
import json
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import re
import time
import subprocess
import platform
import shutil

class ResumeBuilder:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)
        
    def load_profile(self, profile_path: str) -> dict:
        """Load user's base profile including experience and skills"""
        with open(profile_path, 'r') as file:
            return json.load(file)
            
    def parse_job_posting(self, job_description: str) -> dict:
        """Extract key requirements and keywords from job posting"""
        prompt = f"""
        Analyze this job posting and extract:
        1. Required skills
        2. Required experience
        3. Key responsibilities
        4. Industry-specific keywords
        
        Return the response in JSON format.
        
        Job posting:
        {job_description}
        """
        
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        
        return json.loads(response.choices[0].message.content)
        
    def generate_tailored_resume(self, profile: dict, job_requirements: dict) -> str:
        """Generate an ATS-optimized resume based on profile and job requirements"""
        prompt = f"""
        Create an ATS-optimized resume using the candidate's profile and job requirements.
        Follow these strict formatting rules:
        1. Use reverse chronological order for work history
        2. Include both full terms and acronyms for technical terms (e.g., "Project Management Professional (PMP)")
        3. Use standard section headings: "WORK EXPERIENCE", "EDUCATION", "SKILLS"
        4. Format dates as "Month YYYY"
        5. Use bullet points for achievements and responsibilities
        6. Incorporate relevant keywords from the job requirements naturally within experience descriptions
        
        Candidate Profile:
        {json.dumps(profile, indent=2)}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        """
        
        response = self.client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.choices[0].message.content

    def create_docx_resume(self, content: str, output_filename: str) -> str:
        """Convert the resume content to a formatted DOCX file with ATS-friendly formatting"""
        doc = Document()
        
        # Set ATS-friendly margins (1 inch on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Define styles
        styles = doc.styles
        
        # Heading style (14-16pt)
        heading_style = styles.add_style('ATS Heading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Arial'
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Body text style (11-12pt)
        body_style = styles.add_style('ATS Body', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(11)

        # Split content into sections
        sections = re.split(r'\n(?=[A-Z][A-Z\s]+:)', content)
        
        for section in sections:
            if ':' in section:
                title, content = section.split(':', 1)
                
                # Standardize section headings
                title = title.strip().upper()
                if "EXPERIENCE" in title:
                    title = "WORK EXPERIENCE"
                elif "EDUCATION" in title:
                    title = "EDUCATION"
                elif "SKILLS" in title:
                    title = "SKILLS"
                
                # Add heading
                heading = doc.add_paragraph(style='ATS Heading')
                heading.add_run(title)
                
                # Process content based on section type
                content_lines = content.strip().split('\n')
                
                if title == "WORK EXPERIENCE":
                    # Format work experience entries
                    current_entry = []
                    for line in content_lines:
                        if line.strip():
                            if re.match(r'^[A-Za-z\s]+\s+\|\s+', line):  # New job entry
                                if current_entry:
                                    self._format_job_entry(doc, current_entry)
                                current_entry = [line]
                            else:
                                current_entry.append(line)
                    if current_entry:
                        self._format_job_entry(doc, current_entry)
                
                else:
                    # Format other sections
                    for line in content_lines:
                        if line.strip():
                            p = doc.add_paragraph(style='ATS Body')
                            p.add_run(line.strip())
            
            else:
                # Handle header information
                p = doc.add_paragraph(style='ATS Body')
                p.add_run(section.strip())
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Save the document
        docx_path = f"{output_filename}.docx"
        doc.save(docx_path)
        return docx_path

    def _format_job_entry(self, doc, entry_lines):
        """Format a single job entry with proper ATS formatting"""
        # First line contains company and dates
        header = entry_lines[0].split('|')
        company = header[0].strip()
        dates = header[1].strip() if len(header) > 1 else ""
        
        # Format company and dates
        p = doc.add_paragraph(style='ATS Body')
        p.add_run(company).bold = True
        if dates:
            p.add_run(f" | {dates}")
        
        # Format title if present
        if len(entry_lines) > 1:
            p = doc.add_paragraph(style='ATS Body')
            p.add_run(entry_lines[1].strip()).italic = True
        
        # Format bullet points
        for line in entry_lines[2:]:
            if line.strip():
                p = doc.add_paragraph(style='ATS Body')
                p.style.paragraph_format.left_indent = Inches(0.25)
                p.style.paragraph_format.first_line_indent = Inches(-0.25)
                p.add_run("• " + line.strip())

    def flatten_pdf(self, input_pdf: str, output_pdf: str):
        """Flatten PDF with enhanced error handling"""
        print(f"Attempting to flatten PDF: {input_pdf} -> {output_pdf}")
        print(f"Input file exists: {os.path.exists(input_pdf)}")
        
        try:
            reader = PdfReader(input_pdf)
            writer = PdfWriter()

            for page in reader.pages:
                writer.add_page(page)

            for page in writer.pages:
                if '/Annots' in page:
                    del page['/Annots']

            print(f"Writing flattened PDF to: {output_pdf}")
            with open(output_pdf, 'wb') as output_file:
                writer.write(output_file)
            
            print(f"Flattened PDF created: {os.path.exists(output_pdf)}")
            
        except Exception as e:
            print(f"Flattening error: {str(e)}")
            print(f"Error type: {type(e)}")
            raise

    def convert_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """Convert DOCX to PDF using LibreOffice"""
        try:
            if not hasattr(self, 'libreoffice_path'):
                print("ERROR: LibreOffice path not set")
                return False
                
            print("\nDEBUG: Starting PDF conversion using LibreOffice")
            print(f"DOCX path: {os.path.abspath(docx_path)}")
            print(f"PDF path: {os.path.abspath(pdf_path)}")
            
            output_dir = os.path.dirname(os.path.abspath(pdf_path))
            
            cmd = [
                self.libreoffice_path,
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                output_dir,
                docx_path
            ]
            
            print(f"Running command: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"Conversion failed: {result.stderr}")
                return False
                
            expected_pdf = os.path.join(output_dir, 
                                      os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
            
            if os.path.exists(expected_pdf):
                if expected_pdf != pdf_path:
                    os.replace(expected_pdf, pdf_path)
                print(f"PDF created successfully: {pdf_path}")
                return True
            else:
                print(f"PDF not created at expected location: {expected_pdf}")
                return False
                
        except Exception as e:
            print(f"Error during conversion: {str(e)}")
            return False

    def test_libreoffice(self):
        """Test if LibreOffice is properly installed and accessible"""
        try:
            # Check the operating system
            system = platform.system()
            
            # Try multiple possible LibreOffice locations
            if system == "Darwin":  # macOS
                possible_paths = [
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                    '/opt/homebrew/bin/soffice',
                    '/usr/local/bin/soffice'
                ]
            elif system == "Linux":
                possible_paths = [
                    '/usr/bin/soffice',
                    '/usr/bin/libreoffice'
                ]
            else:  # Windows
                possible_paths = [
                    r'C:\Program Files\LibreOffice\program\soffice.exe',
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe'
                ]
            
            # Also check PATH
            soffice_in_path = shutil.which('soffice')
            if soffice_in_path:
                possible_paths.append(soffice_in_path)
                
            print("DEBUG: Checking LibreOffice paths:")
            for path in possible_paths:
                print(f"Checking: {path}")
                if os.path.exists(path):
                    print(f"Found LibreOffice at: {path}")
                    try:
                        result = subprocess.run([path, '--version'], 
                                             capture_output=True, 
                                             text=True,
                                             timeout=5)
                        if result.returncode == 0:
                            print(f"LibreOffice version: {result.stdout.strip()}")
                            self.libreoffice_path = path  # Store the working path
                            return True
                    except subprocess.TimeoutExpired:
                        print(f"Timeout while checking {path}")
                    except Exception as e:
                        print(f"Error checking {path}: {str(e)}")
            
            # If we get here, no working LibreOffice installation was found
            print("\nDEBUG: Checking brew installation:")
            try:
                brew_result = subprocess.run(['brew', 'list', '--cask'], 
                                          capture_output=True, 
                                          text=True)
                if 'libreoffice' in brew_result.stdout:
                    print("LibreOffice is installed via brew but not found in standard locations")
                else:
                    print("LibreOffice not found in brew installations")
            except Exception as e:
                print(f"Error checking brew: {str(e)}")
                
            return False
            
        except Exception as e:
            print(f"Error testing LibreOffice: {str(e)}")
            print("Stack trace:", traceback.format_exc())
            return False

    def test_pdf_conversion(self):
        """Test PDF conversion with a simple document"""
        from docx import Document
        
        # Create a test document
        test_doc = Document()
        test_doc.add_paragraph('Test PDF conversion')
        test_docx = 'test_conversion.docx'
        test_pdf = 'test_conversion.pdf'
        
        # Save test document
        test_doc.save(test_docx)
        
        # Try conversion
        result = self.convert_to_pdf(test_docx, test_pdf)
        
        # Clean up
        if os.path.exists(test_docx):
            os.remove(test_docx)
        if os.path.exists(test_pdf):
            os.remove(test_pdf)
            
        return result

def main():
    # Load environment variables
    load_dotenv()
    api_key = os.getenv('OPENAI_API_KEY')
    
    if not api_key:
        print("Error: OPENAI_API_KEY not found in .env file")
        return
    
    # Initialize resume builder
    resume_builder = ResumeBuilder(api_key)
    
    # Test LibreOffice first
    print("\n🧪 Testing LibreOffice installation...")
    if not resume_builder.test_libreoffice():
        print("❌ LibreOffice not properly installed or accessible")
        print("Please install LibreOffice using:")
        print("brew install --cask libreoffice")
        return
    
    try:
        # Load profile
        profile = resume_builder.load_profile("my_profile.json")
        print("✅ Profile loaded successfully")
        
        # Read job posting from file
        job_posting_file = "job_posting.txt"
        
        if not os.path.exists(job_posting_file):
            print(f"\n❌ Job posting file '{job_posting_file}' not found.")
            print("Please create a file named 'job_posting.txt' with the job description.")
            return
            
        try:
            with open(job_posting_file, 'r', encoding='utf-8') as file:
                job_posting = file.read()
                if not job_posting.strip():
                    print("❌ Job posting file is empty!")
                    return
                print("✅ Job posting loaded successfully")
        except Exception as e:
            print(f"❌ Error reading job posting file: {str(e)}")
            return
        
        print("\n🔍 Analyzing job posting...")
        job_requirements = resume_builder.parse_job_posting(job_posting)
        
        print("\n📝 Generating tailored resume...")
        tailored_resume = resume_builder.generate_tailored_resume(profile, job_requirements)
        
        # Base filename for outputs
        base_filename = "tailored_resume"
        
        # Save as text file
        with open(f"{base_filename}.txt", "w") as f:
            f.write(tailored_resume)
        
        # Create DOCX
        print("\n📎 Creating DOCX version...")
        docx_path = resume_builder.create_docx_resume(tailored_resume, base_filename)
        
        # Convert to PDF
        print("📄 Creating PDF version...")
        pdf_path = f"{base_filename}.pdf"
        
        # Debug: Print current working directory
        print(f"Current working directory: {os.getcwd()}")
        
        # Debug: List files before PDF creation
        print("Files before PDF creation:")
        print(os.listdir())
        
        if resume_builder.convert_to_pdf(docx_path, pdf_path):
            print("✅ PDF created successfully")
            
            # Debug: List files after PDF creation
            print("Files after PDF creation:")
            print(os.listdir())
            
            # Debug: Check if PDF exists and is readable
            if os.path.exists(pdf_path):
                print(f"PDF file exists at: {os.path.abspath(pdf_path)}")
                print(f"PDF file size: {os.path.getsize(pdf_path)} bytes")
            else:
                print(f"PDF file not found at: {os.path.abspath(pdf_path)}")
            
            # Add a small delay to ensure file is fully written
            time.sleep(2)
            
            # Flatten the PDF
            print("🔨 Flattening PDF for ATS optimization...")
            try:
                flattened_pdf = f"{base_filename}_flattened.pdf"
                
                # Debug: Verify source PDF before flattening
                if not os.path.exists(pdf_path):
                    print(f"Source PDF not found before flattening: {pdf_path}")
                    print("Current directory contents:")
                    print(os.listdir())
                
                resume_builder.flatten_pdf(pdf_path, flattened_pdf)
                
                # Debug: Verify flattened PDF was created
                if os.path.exists(flattened_pdf):
                    print(f"Flattened PDF created successfully: {flattened_pdf}")
                    os.replace(flattened_pdf, pdf_path)
                    print("✨ Resume generated successfully!")
                else:
                    print(f"Flattened PDF not created: {flattened_pdf}")
                
            except Exception as flatten_error:
                print(f"⚠️ PDF flattening failed: {str(flatten_error)}")
                print(f"Error type: {type(flatten_error)}")
                print(f"Error details: {flatten_error.__dict__}")
                
                # Debug: Check file permissions
                try:
                    with open(pdf_path, 'rb') as f:
                        print("PDF file is readable")
                except Exception as e:
                    print(f"Cannot read PDF file: {str(e)}")
        else:
            print("⚠️ PDF conversion failed")
            print(f"Files created:")
            print(f"- {base_filename}.txt")
            print(f"- {base_filename}.docx")
            
            print("\nTo create PDF manually:")
            print("1. Open the DOCX file in Microsoft Word/LibreOffice")
            print("2. Save/Export as PDF")
            print("3. Use 'Save as PDF' option for best results")
        
        print("\nGenerated Resume:\n")
        print(tailored_resume)
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()